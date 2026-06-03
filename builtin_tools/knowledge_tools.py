"""
lambdagent.builtin_tools.knowledge_tools — Knowledge base pipeline

ChunkSplitter  λx. split(text, strategy, chunk_size)
OCRTool        λx. ocr(file_path) → extracted text
DocGen         λx. generate(source_md, format, output_path)
KBManager      λx. manage(action, kb_name, ...)
"""
from __future__ import annotations

import json
import math
import os
import re
import subprocess
import hashlib
from typing import Any, Dict, List, Optional


# ════════════════════════════════════════════════════════════
# A34: ChunkSplitter
# ════════════════════════════════════════════════════════════

class ChunkSplitter:
    """Split long documents into retrieval-friendly chunks."""

    @staticmethod
    def split(text: str, strategy: str = "paragraph", chunk_size: int = 500,
              overlap: int = 50) -> List[str]:
        """Split text into chunks.

        Strategies:
          paragraph — split on double newlines
          heading   — split on Markdown headings (# / ## / ###)
          fixed     — fixed size with overlap
          sentence  — split on sentence boundaries
        """
        if strategy == "paragraph":
            return ChunkSplitter._split_paragraph(text, chunk_size)
        elif strategy == "heading":
            return ChunkSplitter._split_heading(text)
        elif strategy == "sentence":
            return ChunkSplitter._split_sentence(text, chunk_size, overlap)
        else:  # fixed
            return ChunkSplitter._split_fixed(text, chunk_size, overlap)

    @staticmethod
    def _split_paragraph(text: str, max_size: int) -> List[str]:
        paragraphs = re.split(r'\n\s*\n', text.strip())
        chunks = []
        current = ""
        for p in paragraphs:
            p = p.strip()
            if not p:
                continue
            if len(current) + len(p) + 2 > max_size and current:
                chunks.append(current.strip())
                current = p
            else:
                current = current + "\n\n" + p if current else p
        if current.strip():
            chunks.append(current.strip())
        return chunks if chunks else [text.strip()]

    @staticmethod
    def _split_heading(text: str) -> List[str]:
        sections = re.split(r'(?=^#{1,3}\s)', text, flags=re.MULTILINE)
        chunks = [s.strip() for s in sections if s.strip()]
        return chunks if chunks else [text.strip()]

    @staticmethod
    def _split_fixed(text: str, size: int, overlap: int) -> List[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = start + size
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start = end - overlap if overlap > 0 else end
        return chunks if chunks else [text.strip()]

    @staticmethod
    def _split_sentence(text: str, max_size: int, overlap: int) -> List[str]:
        sentences = re.split(r'(?<=[.!?。！？\n])\s+', text)
        chunks = []
        current = ""
        for s in sentences:
            if len(current) + len(s) + 1 > max_size and current:
                chunks.append(current.strip())
                # Keep overlap
                words = current.split()
                overlap_text = " ".join(words[-overlap//4:]) if overlap > 0 else ""
                current = overlap_text + " " + s if overlap_text else s
            else:
                current = current + " " + s if current else s
        if current.strip():
            chunks.append(current.strip())
        return chunks if chunks else [text.strip()]


def chunk_split(input_val: Any) -> str:
    """Tool function for ChunkSplitter."""
    params = _parse(input_val)
    text = params.get("text", "")
    if not text:
        file_path = params.get("file_path", "")
        if file_path and os.path.isfile(file_path):
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        else:
            return "[ERROR] text or file_path required"

    strategy = params.get("strategy", "paragraph")
    chunk_size = params.get("chunk_size", 500)
    overlap = params.get("overlap", 50)
    chunks = ChunkSplitter.split(text, strategy, chunk_size, overlap)
    return json.dumps({
        "chunks": len(chunks),
        "strategy": strategy,
        "preview": [c[:100] + "..." if len(c) > 100 else c for c in chunks[:5]],
        "total_chars": sum(len(c) for c in chunks),
    }, ensure_ascii=False, indent=2)


# ════════════════════════════════════════════════════════════
# A35: OCRTool
# ════════════════════════════════════════════════════════════

def _detect_ocr_backend() -> str:
    """Detect available OCR backend."""
    # PaddleOCR
    try:
        import paddleocr
        return "paddle"
    except ImportError:
        pass
    # Tesseract
    try:
        result = subprocess.run(["tesseract", "--version"], capture_output=True, timeout=5)
        if result.returncode == 0:
            return "tesseract"
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return ""


def ocr_extract(input_val: Any) -> str:
    """Extract text from image or scanned PDF via OCR."""
    params = _parse(input_val)
    file_path = params.get("file_path", "")
    if not file_path:
        return "[ERROR] file_path is required"
    file_path = os.path.expanduser(file_path)
    if not os.path.exists(file_path):
        return f"[ERROR] File not found: {file_path}"

    language = params.get("language", "chi_sim+eng")
    backend = _detect_ocr_backend()

    ext = os.path.splitext(file_path)[1].lower()

    if backend == "paddle":
        return _ocr_paddle(file_path, language)
    elif backend == "tesseract":
        return _ocr_tesseract(file_path, language)
    else:
        # Last resort: try tesseract via bash
        try:
            if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"):
                result = subprocess.run(
                    ["tesseract", file_path, "stdout", "-l", language],
                    capture_output=True, text=True, timeout=30,
                )
                if result.returncode == 0:
                    return result.stdout.strip() or "[EMPTY] No text extracted"
        except Exception:
            pass
        return (
            "[ERROR] No OCR backend available. Install one of:\n"
            "  pip install paddleocr paddlepaddle\n"
            "  brew install tesseract (macOS) / apt install tesseract-ocr (Linux)"
        )


def _ocr_paddle(file_path: str, language: str) -> str:
    try:
        from paddleocr import PaddleOCR
        lang = "ch" if "chi" in language else "en"
        ocr = PaddleOCR(use_angle_cls=True, lang=lang, show_log=False)
        results = ocr.ocr(file_path, cls=True)
        lines = []
        if results:
            for page in results:
                if page:
                    for line in page:
                        if line and len(line) >= 2:
                            text = line[1][0] if isinstance(line[1], (list, tuple)) else str(line[1])
                            lines.append(text)
        return "\n".join(lines) if lines else "[EMPTY] No text extracted"
    except Exception as e:
        return f"[OCR_ERROR] PaddleOCR: {e}"


def _ocr_tesseract(file_path: str, language: str) -> str:
    try:
        result = subprocess.run(
            ["tesseract", file_path, "stdout", "-l", language],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0:
            text = result.stdout.strip()
            return text if text else "[EMPTY] No text extracted"
        return f"[OCR_ERROR] tesseract exit {result.returncode}: {result.stderr[:200]}"
    except subprocess.TimeoutExpired:
        return "[TIMEOUT] OCR timed out after 60s"
    except Exception as e:
        return f"[OCR_ERROR] tesseract: {e}"


# ════════════════════════════════════════════════════════════
# A36: DocGen — Document generation
# ════════════════════════════════════════════════════════════

def doc_generate(input_val: Any) -> str:
    """Generate PDF/Word/HTML from Markdown source."""
    params = _parse(input_val)
    source = params.get("source", "")
    fmt = params.get("format", "html")
    output = params.get("output", "")

    if not source:
        return "[ERROR] source (markdown file path or text) is required"

    # Read source
    if os.path.isfile(source):
        with open(source, "r", encoding="utf-8") as f:
            md_text = f.read()
        if not output:
            base = os.path.splitext(source)[0]
            output = f"{base}.{fmt}"
    else:
        md_text = source
        if not output:
            output = f"output.{fmt}"

    output = os.path.expanduser(output)
    os.makedirs(os.path.dirname(output) or ".", exist_ok=True)

    if fmt == "html":
        return _gen_html(md_text, output)
    elif fmt == "pdf":
        return _gen_pdf(md_text, output)
    elif fmt in ("docx", "word"):
        return _gen_docx(md_text, output)
    else:
        return f"[ERROR] Unsupported format: {fmt}. Use html/pdf/docx"


def _gen_html(md_text: str, output: str) -> str:
    """Markdown → HTML."""
    # Try markdown lib
    try:
        import markdown
        html = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
    except ImportError:
        # Minimal conversion
        html = md_text.replace("\n\n", "</p><p>").replace("\n", "<br>")
        html = f"<p>{html}</p>"

    full_html = f"""<!DOCTYPE html>
<html><head>
<meta charset="utf-8">
<style>
body {{ font-family: -apple-system, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background: #f5f5f5; }}
code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
pre {{ background: #f4f4f4; padding: 16px; border-radius: 6px; overflow-x: auto; }}
</style>
</head><body>
{html}
</body></html>"""

    with open(output, "w", encoding="utf-8") as f:
        f.write(full_html)
    return f"[OK] HTML generated: {output} ({len(full_html)} bytes)"


def _gen_pdf(md_text: str, output: str) -> str:
    """Markdown → PDF."""
    # Strategy 1: weasyprint
    try:
        import weasyprint
        html_result = _gen_html(md_text, output + ".tmp.html")
        weasyprint.HTML(filename=output + ".tmp.html").write_pdf(output)
        os.unlink(output + ".tmp.html")
        return f"[OK] PDF generated: {output}"
    except ImportError:
        pass

    # Strategy 2: pandoc via bash
    try:
        import tempfile
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write(md_text)
            tmp_md = f.name
        result = subprocess.run(
            ["pandoc", tmp_md, "-o", output, "--pdf-engine=xelatex"],
            capture_output=True, text=True, timeout=60,
        )
        os.unlink(tmp_md)
        if result.returncode == 0:
            return f"[OK] PDF generated via pandoc: {output}"
        return f"[ERROR] pandoc failed: {result.stderr[:200]}"
    except FileNotFoundError:
        pass

    # Strategy 3: generate HTML as fallback
    html_path = output.replace(".pdf", ".html")
    _gen_html(md_text, html_path)
    return f"[FALLBACK] PDF generation requires weasyprint or pandoc. HTML saved: {html_path}"


def _gen_docx(md_text: str, output: str) -> str:
    """Markdown → Word."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for line in md_text.split("\n"):
            line = line.rstrip()
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line)
        doc.save(output)
        return f"[OK] Word document generated: {output}"
    except ImportError:
        return "[ERROR] python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"[ERROR] docx generation failed: {e}"


# ════════════════════════════════════════════════════════════
# A37: Knowledge Base Manager
# ════════════════════════════════════════════════════════════

KB_BASE = os.path.expanduser("~/.lambdagent/knowledge_bases")


def kb_manage(input_val: Any) -> str:
    """Knowledge base management tool."""
    params = _parse(input_val)
    action = params.get("action", "list")

    if action == "create":
        return _kb_create(params.get("name", ""), params.get("description", ""))
    elif action == "add":
        return _kb_add(params.get("name", ""), params.get("file_path", ""),
                       params.get("pattern", ""), params.get("chunk", True))
    elif action == "search":
        return _kb_search(params.get("name", ""), params.get("query", ""),
                          params.get("top_k", 5))
    elif action == "list":
        return _kb_list()
    elif action == "delete":
        return _kb_delete(params.get("name", ""))
    elif action == "info":
        return _kb_info(params.get("name", ""))
    else:
        return f"[ERROR] Unknown action: {action}. Use: create/add/search/list/delete/info"


def _kb_create(name: str, description: str) -> str:
    if not name:
        return "[ERROR] name is required"
    kb_dir = os.path.join(KB_BASE, name)
    if os.path.exists(kb_dir):
        return f"[ERROR] Knowledge base '{name}' already exists"
    os.makedirs(kb_dir, exist_ok=True)
    meta = {"name": name, "description": description, "documents": [], "total_chunks": 0}
    with open(os.path.join(kb_dir, "meta.json"), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    return f"[OK] Knowledge base '{name}' created at {kb_dir}"


def _kb_add(name: str, file_path: str, pattern: str, chunk: bool) -> str:
    if not name:
        return "[ERROR] name is required"
    kb_dir = os.path.join(KB_BASE, name)
    if not os.path.exists(kb_dir):
        return f"[ERROR] Knowledge base '{name}' not found"

    # Load meta
    meta_path = os.path.join(kb_dir, "meta.json")
    with open(meta_path, "r") as f:
        meta = json.load(f)

    files_to_add = []
    if file_path:
        file_path = os.path.expanduser(file_path)
        if os.path.isdir(file_path):
            import glob
            pat = pattern or "**/*.*"
            files_to_add = [f for f in glob.glob(os.path.join(file_path, pat), recursive=True)
                           if os.path.isfile(f)]
        elif os.path.isfile(file_path):
            files_to_add = [file_path]
        else:
            return f"[ERROR] Path not found: {file_path}"
    else:
        return "[ERROR] file_path is required"

    # Load or create store
    import sys
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
    from lambdagent.rag import SimpleVectorStore

    store_path = os.path.join(kb_dir, "store.json")
    store = SimpleVectorStore()

    # Load existing store
    if os.path.exists(store_path):
        with open(store_path, "r") as f:
            store_data = json.load(f)
        from lambdagent.rag import Document
        for d in store_data.get("documents", []):
            store.documents.append(Document(content=d["content"], metadata=d.get("metadata", {}), doc_id=d.get("doc_id", "")))

    added = 0
    for fp in files_to_add:
        try:
            with open(fp, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
        except Exception:
            continue

        if chunk and len(text) > 500:
            chunks = ChunkSplitter.split(text, "paragraph", 500, 50)
        else:
            chunks = [text]

        for i, c in enumerate(chunks):
            meta_info = {"source": fp, "chunk_index": i}
            store.add(c, meta_info)
            added += 1

        meta["documents"].append({"path": fp, "chunks": len(chunks)})

    meta["total_chunks"] = len(store.documents)

    # Save store
    store_data = {
        "documents": [
            {"content": d.content, "metadata": d.metadata, "doc_id": d.doc_id}
            for d in store.documents
        ]
    }
    with open(store_path, "w", encoding="utf-8") as f:
        json.dump(store_data, f, ensure_ascii=False)

    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)

    return f"[OK] Added {added} chunks from {len(files_to_add)} file(s) to '{name}'. Total: {meta['total_chunks']} chunks"


def _kb_search(name: str, query: str, top_k: int) -> str:
    if not name or not query:
        return "[ERROR] name and query are required"
    kb_dir = os.path.join(KB_BASE, name)
    store_path = os.path.join(kb_dir, "store.json")
    if not os.path.exists(store_path):
        return f"[ERROR] Knowledge base '{name}' not found or empty"

    from lambdagent.rag import SimpleVectorStore, Document

    store = SimpleVectorStore()
    with open(store_path, "r") as f:
        store_data = json.load(f)
    for d in store_data.get("documents", []):
        store.documents.append(Document(content=d["content"], metadata=d.get("metadata", {}), doc_id=d.get("doc_id", "")))

    results = store.search(query, top_k=top_k)
    if not results:
        return f"[NO_MATCH] No results for '{query}' in '{name}'"

    lines = []
    for r in results:
        source = r.document.metadata.get("source", "?")
        lines.append(f"[Score: {r.score:.3f}] (from: {os.path.basename(source)})")
        lines.append(r.document.content[:300])
        lines.append("")
    return "\n".join(lines)


def _kb_list() -> str:
    if not os.path.exists(KB_BASE):
        return "[EMPTY] No knowledge bases found"
    kbs = []
    for name in sorted(os.listdir(KB_BASE)):
        meta_path = os.path.join(KB_BASE, name, "meta.json")
        if os.path.exists(meta_path):
            with open(meta_path) as f:
                meta = json.load(f)
            kbs.append(f"  {name}: {meta.get('total_chunks', 0)} chunks, {len(meta.get('documents', []))} files — {meta.get('description', '')}")
    if not kbs:
        return "[EMPTY] No knowledge bases found"
    return "Knowledge bases:\n" + "\n".join(kbs)


def _kb_delete(name: str) -> str:
    if not name:
        return "[ERROR] name is required"
    kb_dir = os.path.join(KB_BASE, name)
    if not os.path.exists(kb_dir):
        return f"[ERROR] Knowledge base '{name}' not found"
    import shutil
    shutil.rmtree(kb_dir)
    return f"[OK] Knowledge base '{name}' deleted"


def _kb_info(name: str) -> str:
    if not name:
        return "[ERROR] name is required"
    meta_path = os.path.join(KB_BASE, name, "meta.json")
    if not os.path.exists(meta_path):
        return f"[ERROR] Knowledge base '{name}' not found"
    with open(meta_path) as f:
        meta = json.load(f)
    lines = [
        f"Knowledge Base: {name}",
        f"  Description: {meta.get('description', '')}",
        f"  Total chunks: {meta.get('total_chunks', 0)}",
        f"  Documents ({len(meta.get('documents', []))}):",
    ]
    for doc in meta.get("documents", [])[:10]:
        lines.append(f"    - {doc.get('path', '?')} ({doc.get('chunks', 0)} chunks)")
    return "\n".join(lines)


# ════════════════════════════════════════════════════════════
# Shared
# ════════════════════════════════════════════════════════════

def _parse(input_val: Any) -> dict:
    if isinstance(input_val, dict):
        return input_val
    if isinstance(input_val, str):
        try:
            return json.loads(input_val)
        except (json.JSONDecodeError, ValueError):
            return {"text": input_val} if len(input_val) > 50 else {"query": input_val}
    return {}
